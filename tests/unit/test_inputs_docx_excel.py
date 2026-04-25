from __future__ import annotations

import base64
import json
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from ai_qahelper.inputs import parse_requirements
from ai_qahelper.llm_client import LlmClient
from ai_qahelper.models import AppConfig

_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def test_parse_docx_extracts_paragraphs_headings_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "requirements.docx"
    doc = Document()
    doc.add_heading("Оплата заказа", level=1)
    doc.add_paragraph("Пользователь может оплатить заказ картой.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Поле"
    table.cell(0, 1).text = "Правило"
    table.cell(1, 0).text = "Email"
    table.cell(1, 1).text = "Обязательное"
    doc.save(docx_path)

    items = parse_requirements([str(docx_path)])

    content = items[0].content
    assert "# DOCX requirements: requirements.docx" in content
    assert "### Оплата заказа" in content
    assert "Пользователь может оплатить заказ картой." in content
    assert "Table 1:" in content
    assert "| Поле | Правило |" in content
    assert "| Email | Обязательное |" in content


def test_parse_docx_with_image_uses_vision_and_writes_coverage(monkeypatch, tmp_path: Path) -> None:
    docx_path = tmp_path / "visual.docx"
    coverage_path = tmp_path / "input-coverage-report.json"
    doc = Document()
    doc.add_picture(BytesIO(_PNG_1X1))
    doc.save(docx_path)

    def fake_describe(self, images, *, images_per_batch=2, max_output_tokens=4096):  # noqa: ANN001
        assert len(images) == 1
        return "Image 001-image1.png: На экране видна кнопка «Оплатить»."

    monkeypatch.setattr(LlmClient, "describe_images_for_requirements", fake_describe)
    cfg = AppConfig.model_validate({"llm": {"api_key": "test-key"}})

    items = parse_requirements([str(docx_path)], cfg, coverage_report_path=coverage_path, session_dir=tmp_path)

    assert "кнопка «Оплатить»" in items[0].content
    coverage = json.loads(coverage_path.read_text(encoding="utf-8"))
    assert coverage["type"] == "docx"
    assert coverage["images_found"] == 1
    assert coverage["images_analyzed"] == 1
    assert coverage["images_skipped"] == 0


def test_parse_docx_image_warns_when_vision_disabled(tmp_path: Path) -> None:
    docx_path = tmp_path / "visual-disabled.docx"
    coverage_path = tmp_path / "input-coverage-report.json"
    doc = Document()
    doc.add_picture(BytesIO(_PNG_1X1))
    doc.save(docx_path)
    cfg = AppConfig.model_validate({"llm": {"api_key": "test-key", "docx_vision": False}})

    items = parse_requirements([str(docx_path)], cfg, coverage_report_path=coverage_path, session_dir=tmp_path)

    assert "docx_vision is disabled" in items[0].content
    coverage = json.loads(coverage_path.read_text(encoding="utf-8"))
    assert coverage["images_found"] == 1
    assert coverage["images_analyzed"] == 0
    assert coverage["images_skipped"] == 1
    assert "Visual requirements were not analyzed" in coverage["warnings"][0]
    assert "Risk: incomplete requirements coverage" in coverage["warnings"][0]
    assert "Risk: incomplete requirements coverage" in items[0].content


def test_parse_xlsx_extracts_sheets_rows_and_cells(tmp_path: Path) -> None:
    xlsx_path = tmp_path / "requirements.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Auth"
    sheet.append(["Requirement", "Expected"])
    sheet.append(["Login with valid user", "Dashboard opens"])
    second = workbook.create_sheet("Payments")
    second.append(["Payment", "Status"])
    second.append(["Card", "Approved"])
    workbook.save(xlsx_path)

    items = parse_requirements([str(xlsx_path)])

    content = items[0].content
    assert "# Excel requirements: requirements.xlsx" in content
    assert "## Sheet: Auth" in content
    assert "| Requirement | Expected |" in content
    assert "| Login with valid user | Dashboard opens |" in content
    assert "## Sheet: Payments" in content
    assert "| Card | Approved |" in content


def test_parse_xlsx_truncation_warning(monkeypatch, tmp_path: Path) -> None:
    xlsx_path = tmp_path / "large.xlsx"
    coverage_path = tmp_path / "input-coverage-report.json"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["A", "B"])
    sheet.append(["1", "2"])
    sheet.append(["3", "4"])
    workbook.save(xlsx_path)
    monkeypatch.setattr("ai_qahelper.inputs.MAX_EXCEL_ROWS_PER_SHEET", 2)

    items = parse_requirements([str(xlsx_path)], coverage_report_path=coverage_path)

    assert "Excel file was truncated" in items[0].content
    coverage = json.loads(coverage_path.read_text(encoding="utf-8"))
    assert coverage["truncated"] is True
    assert coverage["warnings"]
