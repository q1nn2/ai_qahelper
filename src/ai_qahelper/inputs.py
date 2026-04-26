from __future__ import annotations

import json
import os
import tempfile
import zipfile
from pathlib import Path
from typing import Any

import httpx
import pandas as pd
from openpyxl import load_workbook
from pypdf import PdfReader

from ai_qahelper.llm_client import LlmClient
from ai_qahelper.models import AppConfig, DesignModel, DesignNode, RequirementItem
from ai_qahelper.pdf_vision import build_pdf_requirement_content

MAX_EXCEL_SHEETS = 10
MAX_EXCEL_ROWS_PER_SHEET = 500
MAX_EXCEL_COLUMNS_PER_SHEET = 50
_DOCX_VISUAL_COVERAGE_RISK = "Risk: incomplete requirements coverage because visual requirements were not analyzed."


def parse_requirements(
    paths: list[str],
    app_cfg: AppConfig | None = None,
    *,
    coverage_report_path: Path | None = None,
    session_dir: Path | None = None,
) -> list[RequirementItem]:
    items: list[RequirementItem] = []
    llm: LlmClient | None = None
    if app_cfg and app_cfg.pdf_vision and any(Path(p).suffix.lower() == ".pdf" for p in paths):
        llm = LlmClient(app_cfg.llm)
    coverage_entries: list[dict[str, Any]] = []

    for path in paths:
        p = Path(path)
        if not p.is_file():
            resolved = p.resolve()
            raise FileNotFoundError(
                f"Requirement file not found: {path!r} (resolved: {resolved}, cwd: {Path.cwd()}). "
                "Check the path and filename (on Windows avoid accidental .pdf.pdf if extensions are hidden)."
            )
        suffix = p.suffix.lower()
        if suffix == ".pdf" and app_cfg is not None and llm is not None:
            content = build_pdf_requirement_content(
                p,
                llm,
                app_cfg.llm,
                pdf_vision=app_cfg.pdf_vision,
            )
            coverage_entries.append(_coverage_base(p, "pdf"))
        else:
            content, coverage = _read_requirement_file(p, app_cfg=app_cfg, session_dir=session_dir)
            coverage_entries.append(coverage)
        items.append(RequirementItem(source=str(p), content=content))
    if coverage_report_path is not None:
        _write_input_coverage_report(coverage_report_path, coverage_entries)
    return items


def _read_requirement_file(
    path: Path,
    *,
    app_cfg: AppConfig | None = None,
    session_dir: Path | None = None,
) -> tuple[str, dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf_requirement(path), _coverage_base(path, "pdf")
    if suffix == ".docx":
        return _read_docx_requirement(path, app_cfg=app_cfg, session_dir=session_dir)
    if suffix == ".xlsx":
        return _read_xlsx_requirement(path)
    if suffix == ".xls":
        return _read_xls_requirement(path)
    return path.read_text(encoding="utf-8"), _coverage_base(path, suffix.lstrip(".") or "text")


def _read_pdf_requirement(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    text = "\n".join(chunks).strip()
    if not text:
        return "PDF parsed but no text content was extracted."
    return text


def _read_docx_requirement(
    path: Path,
    *,
    app_cfg: AppConfig | None = None,
    session_dir: Path | None = None,
) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to read .docx requirement files.") from exc

    doc = Document(str(path))
    paragraphs: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            paragraphs.append(f"### {text}")
        elif "list" in style_name:
            paragraphs.append(f"- {text}")
        else:
            paragraphs.append(text)

    table_blocks: list[str] = []
    for idx, table in enumerate(doc.tables, start=1):
        rows = [[_clean_cell_text(cell.text) for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(cell for cell in row)]
        if rows:
            table_blocks.append(f"Table {idx}:\n{_markdown_table(rows)}")

    images = _extract_docx_images(path, session_dir=session_dir)
    coverage = {
        **_coverage_base(path, "docx"),
        "paragraphs_found": len(paragraphs),
        "tables_found": len(table_blocks),
        "images_found": len(images),
        "images_analyzed": 0,
        "images_skipped": 0,
        "warnings": [],
    }
    warnings: list[str] = coverage["warnings"]
    image_descriptions = _describe_docx_images(path, images, app_cfg=app_cfg, coverage=coverage)

    parts = [f"# DOCX requirements: {path.name}", "", "## Paragraphs"]
    if paragraphs:
        parts.append("\n".join(paragraphs))
    else:
        parts.append("DOCX parsed but no text content was extracted.")

    parts.extend(["", "## Tables"])
    parts.append("\n\n".join(table_blocks) if table_blocks else "No tables found.")

    parts.extend(["", "## Images / Visual requirements"])
    if image_descriptions:
        parts.append(image_descriptions)
    elif images:
        parts.append("Images were found but no visual descriptions were produced.")
    else:
        parts.append("No images found.")

    if warnings:
        parts.extend(["", "## Warnings", "\n".join(f"- {warning}" for warning in warnings)])

    return "\n".join(parts).strip(), coverage


def _extract_docx_images(path: Path, *, session_dir: Path | None) -> list[tuple[str, bytes]]:
    images: list[tuple[str, bytes]] = []
    output_dir = (session_dir / "docx-images" / path.stem) if session_dir else Path(tempfile.mkdtemp(prefix="ai_qahelper_docx_images_"))
    with zipfile.ZipFile(path) as zf:
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/") and not name.endswith("/"))
        for idx, name in enumerate(media_names, start=1):
            data = zf.read(name)
            filename = f"{idx:03d}-{Path(name).name}"
            output_dir.mkdir(parents=True, exist_ok=True)
            (output_dir / filename).write_bytes(data)
            images.append((filename, data))
    return images


def _describe_docx_images(
    path: Path,
    images: list[tuple[str, bytes]],
    *,
    app_cfg: AppConfig | None,
    coverage: dict[str, Any],
) -> str:
    if not images:
        return ""

    warnings: list[str] = coverage["warnings"]
    llm_cfg = app_cfg.llm if app_cfg is not None else None
    if llm_cfg is None or not llm_cfg.docx_vision:
        coverage["images_skipped"] = len(images)
        warning = (
            f"DOCX contains {len(images)} images, but docx_vision is disabled. "
            f"Visual requirements were not analyzed. {_DOCX_VISUAL_COVERAGE_RISK}"
        )
        warnings.append(warning)
        return ""

    max_images = max(0, llm_cfg.docx_vision_max_images)
    to_analyze = images[:max_images]
    skipped = len(images) - len(to_analyze)
    if skipped > 0:
        coverage["images_skipped"] = skipped
        warnings.append(
            f"DOCX contains {len(images)} images; only first {max_images} were analyzed, {skipped} skipped. "
            f"{_DOCX_VISUAL_COVERAGE_RISK}"
        )
    if not to_analyze:
        return ""

    try:
        llm = LlmClient(llm_cfg)
        visual = llm.describe_images_for_requirements(
            to_analyze,
            images_per_batch=max(1, llm_cfg.docx_vision_images_per_request),
            max_output_tokens=max(512, llm_cfg.docx_vision_max_output_tokens),
        )
    except Exception as exc:  # noqa: BLE001 - input parsing must not fail only because vision failed
        coverage["images_skipped"] += len(to_analyze)
        warning = (
            f"DOCX contains images, but vision analysis failed for {path.name}: {exc}. "
            f"{_DOCX_VISUAL_COVERAGE_RISK}"
        )
        warnings.append(warning)
        return ""

    coverage["images_analyzed"] = len(to_analyze)
    lines: list[str] = []
    if visual.strip():
        lines.append(visual.strip())
    else:
        coverage["images_skipped"] += len(to_analyze)
        coverage["images_analyzed"] = 0
        warnings.append(
            "DOCX image vision returned an empty description. "
            f"{_DOCX_VISUAL_COVERAGE_RISK}"
        )
    for filename, _ in images[max_images:]:
        lines.append(f"Image {filename}: skipped because docx_vision_max_images limit was reached.")
    return "\n\n".join(lines)


def _read_xlsx_requirement(path: Path) -> tuple[str, dict[str, Any]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet_names = workbook.sheetnames
        analyzed_names = sheet_names[:MAX_EXCEL_SHEETS]
        coverage = {
            **_coverage_base(path, "xlsx"),
            "sheets_found": len(sheet_names),
            "sheets_analyzed": len(analyzed_names),
            "rows_analyzed": 0,
            "cells_analyzed": 0,
            "truncated": len(sheet_names) > MAX_EXCEL_SHEETS,
            "warnings": [],
        }
        blocks: list[str] = [f"# Excel requirements: {path.name}"]
        for sheet_name in analyzed_names:
            sheet = workbook[sheet_name]
            rows = _extract_excel_rows(sheet.iter_rows(values_only=True), coverage)
            blocks.extend(["", f"## Sheet: {sheet_name}"])
            blocks.append(_markdown_table(rows) if rows else "No non-empty cells found.")
        _add_excel_truncation_warning(coverage)
        if coverage["rows_analyzed"] == 0:
            blocks.append("\nExcel parsed but no non-empty cells were found.")
        if coverage["warnings"]:
            blocks.extend(["", "## Warnings", "\n".join(f"- {w}" for w in coverage["warnings"])])
        return "\n".join(blocks).strip(), coverage
    finally:
        workbook.close()


def _read_xls_requirement(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        sheets = pd.read_excel(path, sheet_name=None, header=None, engine=None)
    except ImportError as exc:
        raise RuntimeError(".xls is not supported without xlrd. Please save file as .xlsx.") from exc
    except ValueError as exc:
        if "xlrd" in str(exc).lower():
            raise RuntimeError(".xls is not supported without xlrd. Please save file as .xlsx.") from exc
        raise

    coverage = {
        **_coverage_base(path, "xls"),
        "sheets_found": len(sheets),
        "sheets_analyzed": min(len(sheets), MAX_EXCEL_SHEETS),
        "rows_analyzed": 0,
        "cells_analyzed": 0,
        "truncated": len(sheets) > MAX_EXCEL_SHEETS,
        "warnings": [],
    }
    blocks: list[str] = [f"# Excel requirements: {path.name}"]
    for sheet_name, frame in list(sheets.items())[:MAX_EXCEL_SHEETS]:
        rows = _extract_excel_rows(frame.itertuples(index=False, name=None), coverage)
        blocks.extend(["", f"## Sheet: {sheet_name}"])
        blocks.append(_markdown_table(rows) if rows else "No non-empty cells found.")
    _add_excel_truncation_warning(coverage)
    if coverage["rows_analyzed"] == 0:
        blocks.append("\nExcel parsed but no non-empty cells were found.")
    if coverage["warnings"]:
        blocks.extend(["", "## Warnings", "\n".join(f"- {w}" for w in coverage["warnings"])])
    return "\n".join(blocks).strip(), coverage


def _extract_excel_rows(rows_iter: Any, coverage: dict[str, Any]) -> list[list[str]]:
    out: list[list[str]] = []
    for raw_idx, raw_row in enumerate(rows_iter, start=1):
        if raw_idx > MAX_EXCEL_ROWS_PER_SHEET:
            coverage["truncated"] = True
            break
        raw_values = list(raw_row)
        row = [_format_cell(value) for value in raw_values[:MAX_EXCEL_COLUMNS_PER_SHEET]]
        if len(raw_values) > MAX_EXCEL_COLUMNS_PER_SHEET:
            coverage["truncated"] = True
        if not any(row):
            continue
        coverage["rows_analyzed"] += 1
        coverage["cells_analyzed"] += sum(1 for cell in row if cell)
        out.append(row)
    if len(out) > 1 and _looks_like_header(out[0]):
        return out
    width = max((len(row) for row in out), default=0)
    return [[f"Column {idx}" for idx in range(1, width + 1)], *out] if out else []


def _add_excel_truncation_warning(coverage: dict[str, Any]) -> None:
    if coverage["truncated"]:
        coverage["warnings"].append(
            "Excel file was truncated: only first 10 sheets and 500 rows per sheet were analyzed."
        )


def _looks_like_header(row: list[str]) -> bool:
    if not row or not any(row):
        return False
    non_empty = [cell for cell in row if cell]
    return len(non_empty) >= 2 and all(not _looks_numeric(cell) for cell in non_empty)


def _looks_numeric(value: str) -> bool:
    try:
        float(value.replace(",", "."))
    except ValueError:
        return False
    return True


def _format_cell(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()


def _clean_cell_text(value: str) -> str:
    return " ".join(line.strip() for line in value.splitlines() if line.strip())


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [(row + [""] * width)[:width] for row in rows]
    header = [_escape_markdown_cell(cell or f"Column {idx}") for idx, cell in enumerate(normalized[0], start=1)]
    body = normalized[1:] or [[""] * width]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(_escape_markdown_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def _escape_markdown_cell(value: str) -> str:
    return (value or "").replace("\n", "<br>").replace("|", "\\|")


def _coverage_base(path: Path, source_type: str) -> dict[str, Any]:
    return {"source": path.name, "type": source_type, "warnings": []}


def _write_input_coverage_report(path: Path, entries: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if len(entries) == 1:
        payload: dict[str, Any] = entries[0]
    else:
        payload = {
            "sources": entries,
            "summary": {
                "files_found": len(entries),
                "warnings": [warning for entry in entries for warning in entry.get("warnings", [])],
            },
        }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_requirement_url(url: str, timeout_s: int = 20) -> RequirementItem:
    with httpx.Client(timeout=timeout_s) as client:
        resp = client.get(url)
        resp.raise_for_status()
    return RequirementItem(source=url, content=resp.text)


def _collect_nodes(raw_node: dict, max_depth: int = 3, depth: int = 0) -> DesignNode:
    node = DesignNode(
        id=raw_node.get("id", "unknown"),
        name=raw_node.get("name", "unnamed"),
        text=raw_node.get("characters"),
        node_type=raw_node.get("type"),
    )
    if depth >= max_depth:
        return node
    for child in raw_node.get("children", [])[:50]:
        node.children.append(_collect_nodes(child, max_depth=max_depth, depth=depth + 1))
    return node


def _summarize_design(model: DesignModel) -> tuple[int, int]:
    nodes_total = 0
    text_nodes = 0
    stack = list(model.nodes)
    while stack:
        node = stack.pop()
        nodes_total += 1
        if node.text:
            text_nodes += 1
        stack.extend(node.children)
    return nodes_total, text_nodes


def ingest_figma(file_key: str, node_ids: list[str] | None = None) -> DesignModel:
    token = os.getenv("FIGMA_API_TOKEN")
    if not token:
        return DesignModel(file_key=file_key, warnings=["Missing FIGMA_API_TOKEN; skipped Figma ingestion"])

    headers = {"X-Figma-Token": token}
    with httpx.Client(timeout=30) as client:
        file_resp = client.get(f"https://api.figma.com/v1/files/{file_key}", headers=headers)
        if file_resp.status_code >= 400:
            return DesignModel(
                file_key=file_key,
                warnings=[f"Figma file fetch failed: HTTP {file_resp.status_code}"],
            )
        file_json = file_resp.json()
        model = DesignModel(file_key=file_key, file_name=file_json.get("name"))

        doc = file_json.get("document", {})
        if doc:
            model.nodes.append(_collect_nodes(doc, max_depth=5))
        else:
            model.warnings.append("Figma file has empty document tree")

        if node_ids:
            ids = ",".join(node_ids)
            nodes_resp = client.get(
                f"https://api.figma.com/v1/files/{file_key}/nodes",
                headers=headers,
                params={"ids": ids},
            )
            if nodes_resp.status_code < 400:
                nodes_payload = nodes_resp.json().get("nodes", {})
                for node_id, body in nodes_payload.items():
                    document = body.get("document")
                    if document:
                        model.nodes.append(_collect_nodes(document, max_depth=5))
            else:
                model.warnings.append(f"Figma nodes fetch failed: HTTP {nodes_resp.status_code}")
        nodes_total, text_nodes = _summarize_design(model)
        if nodes_total == 0:
            model.warnings.append("No Figma nodes were collected")
        if text_nodes == 0:
            model.warnings.append("No textual labels found in collected Figma nodes")
        return model
